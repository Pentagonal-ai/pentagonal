#!/usr/bin/env python3
"""Generate security audit Word document for arcade.somethingdumb.xyz"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)

# ── Styles ──
styles = doc.styles

def set_font(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1, color=(0,0,0)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(22)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(15)
        run.bold = True
    elif level == 3:
        run.font.size = Pt(12)
        run.bold = True
    run.font.color.rgb = RGBColor(*color)
    return p

def add_body(text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(20, 20, 20)
    # Shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def severity_badge(sev):
    badges = {
        'HIGH': ('🔴 HIGH', (180, 0, 0)),
        'MEDIUM': ('🟡 MEDIUM', (160, 100, 0)),
        'LOW': ('🟢 LOW', (0, 100, 0)),
        'INFO': ('ℹ INFO', (0, 70, 140)),
    }
    return badges.get(sev, (sev, (0,0,0)))

# ═══════════════════════════════════════
# COVER
# ═══════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SECURITY AUDIT REPORT')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(15, 15, 15)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('arcade.somethingdumb.xyz')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 200)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Assessment Type: Black-Box Web + API Security Review')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Scope: Full web application, API endpoints, blockchain integration')
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# ═══════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════
doc.add_page_break()

add_heading('1. Executive Summary', 1)

add_body(
    'arcade.somethingdumb.xyz is a Solana-based blockchain arcade platform where players submit '
    'high scores as on-chain memo transactions to compete for weekly $DUMB token prizes. '
    'The application is built on Next.js/Vercel and integrates directly with game logic, '
    'wallet signing, and Solana RPC endpoints.'
)

add_body(
    'The attack surface is significant because real money is on the line — the treasury wallet '
    'holds weekly prize pool funds. The core integrity mechanism (requiring signed Solana transactions '
    'for score submission) is sound in concept, but several implementation weaknesses could allow an '
    'attacker to inflate their score standing, drain wallet sessions via XSS, or observe prize winner '
    'selection state to time attacks.'
)

p = doc.add_paragraph()
run = p.add_run('Overall Risk Rating: ')
run.font.size = Pt(11)
run.bold = True
run = p.add_run('MEDIUM-HIGH')
run.font.size = Pt(11)
run.bold = True
run.font.color.rgb = RGBColor(180, 80, 0)

# ═══════════════════════════════════════
# FINDING SUMMARY TABLE
# ═══════════════════════════════════════
add_heading('2. Finding Summary', 1)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr = table.rows[0].cells
for i, text in enumerate(['#', 'Finding', 'Severity', 'Category']):
    p = hdr[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)

findings_summary = [
    ('1', 'Weak Content Security Policy (unsafe-inline + unsafe-eval)', 'HIGH', 'XSS / Client Security'),
    ('2', 'Score Submission Signature Validation Unknown', 'HIGH', 'Score Integrity / Fraud'),
    ('3', 'Wildcard CORS on Main Page', 'MEDIUM', 'Cross-Origin / CSRF'),
    ('4', 'Missing X-Frame-Options / Clickjacking', 'MEDIUM', 'UI Redress'),
    ('5', 'Lucky Winner Endpoint Exposes Prize State', 'MEDIUM', 'Information Disclosure'),
    ('6', 'Missing HSTS Header', 'LOW', 'Transport Security'),
    ('7', 'Treasury Wallet & Contract Hardcoded in Client', 'LOW', 'Information Disclosure'),
    ('8', 'Leaderboard 503 Under Load', 'LOW', 'Reliability / Race Condition'),
    ('9', 'Balance Endpoints Unauthenticated & Unrated', 'LOW', 'Privacy / Rate Limiting'),
]

sev_colors = {
    'HIGH': RGBColor(180, 0, 0),
    'MEDIUM': RGBColor(160, 100, 0),
    'LOW': RGBColor(0, 120, 0),
}

for num, title, sev, cat in findings_summary:
    row = table.add_row().cells
    row[0].paragraphs[0].add_run(num).font.size = Pt(9)
    row[1].paragraphs[0].add_run(title).font.size = Pt(9)
    r = row[2].paragraphs[0].add_run(sev)
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = sev_colors.get(sev, RGBColor(0,0,0))
    row[3].paragraphs[0].add_run(cat).font.size = Pt(9)

doc.add_paragraph()

# ═══════════════════════════════════════
# DETAILED FINDINGS
# ═══════════════════════════════════════
doc.add_page_break()
add_heading('3. Detailed Findings', 1)

# ── Finding 1 ──
add_heading('Finding 1: Weak Content Security Policy', 2, color=(180,0,0))
p = doc.add_paragraph()
r = p.add_run('Severity: HIGH  |  Category: XSS / Client Security  |  CVSS: 7.5')
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RGBColor(180, 0, 0)

add_body('Where: All pages via response headers', italic=True)
add_body(
    'The application\'s Content-Security-Policy includes both \'unsafe-inline\' and \'unsafe-eval\' '
    'directives for scripts. These directives void the primary XSS defense layer — their entire purpose '
    'is to prevent execution of injected scripts, and both override it.'
)
add_body('Impact:', bold=True)
add_body(
    'If any page reflects user-supplied input (wallet addresses, usernames, leaderboard entries) without '
    'proper encoding, an attacker can execute arbitrary JavaScript in the victim\'s browser session. '
    'In a wallet-connected application this is critical:'
)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Exfiltrate session data and wallet adapter state').font.size = Pt(10)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Manipulate wallet popup approvals via DOM injection (Phantom/Backpack)').font.size = Pt(10)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Silently redirect score transactions to an attacker-controlled wallet').font.size = Pt(10)

add_body('Remediation:', bold=True)
add_code(
    "Content-Security-Policy:\n"
    "  default-src 'self';\n"
    "  script-src 'self' 'nonce-{server-generated}';\n"
    "  connect-src 'self' https://api.mainnet-beta.solana.com;\n"
    "  img-src 'self' data: https:;\n"
    "  style-src 'self' 'unsafe-inline';"
)
add_body('Remove unsafe-eval — Next.js 15 supports nonce-based CSP. Never use unsafe-inline in production.', italic=True)

doc.add_paragraph()

# ── Finding 2 ──
add_heading('Finding 2: Score Submission Signature Validation Unknown', 2, color=(180,0,0))
p = doc.add_paragraph()
r = p.add_run('Severity: HIGH  |  Category: Score Integrity / Fraud  |  CVSS: 8.1')
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RGBColor(180, 0, 0)

add_body('Where: /api/send-tx', italic=True)
add_body(
    'The score submission endpoint accepts serialized Solana transactions and correctly requires a '
    'properly-formed transaction object. However it is unknown whether the server validates:'
)
p = doc.add_paragraph(style='List Number')
p.add_run('That the transaction signer matches the wallet claiming the score').font.size = Pt(10)
p = doc.add_paragraph(style='List Number')
p.add_run('That the score embedded in the memo is within plausible game bounds').font.size = Pt(10)
p = doc.add_paragraph(style='List Number')
p.add_run('That the blockhash is recent (replay prevention window)').font.size = Pt(10)
p = doc.add_paragraph(style='List Number')
p.add_run('That memo content is unique per session (cannot be replayed)').font.size = Pt(10)

add_body('Proof-of-Concept Attack:', bold=True)
add_code(
    "// 1. Play legitimately and intercept the signed transaction\n"
    "const tx = interceptedSerializedTx; // from one real game session\n\n"
    "// 2. Replay the SAME transaction N times\n"
    "for (let i = 0; i < 100; i++) {\n"
    "  await fetch('/api/send-tx', {\n"
    "    method: 'POST',\n"
    "    body: JSON.stringify({ tx })\n"
    "  });\n"
    "}\n"
    "// Result: 100 leaderboard entries from one real gameplay session"
)
add_body('If successful: attacker climbs to #1 leaderboard position and claims weekly prize pool.', italic=True)

add_body('Remediation:', bold=True)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Issue server-side HMAC session token at game start, embed in memo').font.size = Pt(10)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Server validates HMAC on /api/send-tx before accepting score').font.size = Pt(10)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Enforce blockhash epoch window: reject transactions older than ~2 minutes').font.size = Pt(10)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Add score sanity bounds: reject memos claiming scores > 3σ above known human maximums').font.size = Pt(10)

doc.add_paragraph()

# ── Finding 3 ──
add_heading('Finding 3: Wildcard CORS on Main Page', 2, color=(160,100,0))
p = doc.add_paragraph()
r = p.add_run('Severity: MEDIUM  |  Category: Cross-Origin / CSRF  |  CVSS: 5.3')
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RGBColor(160, 100, 0)

add_body('Where: / and API routes  |  Header observed: access-control-allow-origin: *', italic=True)
add_body(
    'The wildcard CORS header allows any origin to make credentialed cross-origin requests and read '
    'responses. Combined with the XSS finding, a malicious site can silently probe a victim\'s '
    'wallet state or credit balance:'
)
add_code(
    "// From evil.com — works due to wildcard CORS\n"
    "const r = await fetch(\n"
    "  'https://arcade.somethingdumb.xyz/api/balance?wallet=' + victim\n"
    ");\n"
    "const data = await r.json(); // leaks victim's credit balance"
)
add_body('Remediation: Lock CORS to https://arcade.somethingdumb.xyz only. Verify Origin header server-side for wallet-related endpoints.')

doc.add_paragraph()

# ── Finding 4 ──
add_heading('Finding 4: Missing X-Frame-Options / Clickjacking', 2, color=(160,100,0))
p = doc.add_paragraph()
r = p.add_run('Severity: MEDIUM  |  Category: UI Redress  |  CVSS: 5.4')
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RGBColor(160, 100, 0)

add_body('Where: All pages  |  Missing headers: X-Frame-Options, frame-ancestors CSP', italic=True)
add_body(
    'Without clickjacking protection, an attacker can iframe the game inside a malicious page '
    'and overlay invisible UI elements. Users primed to click "approve" in their Solana wallet '
    'can be tricked into approving transactions they didn\'t intend — particularly dangerous '
    'during the score submission flow.'
)
add_body('Remediation:')
add_code(
    "X-Frame-Options: DENY\n"
    "Content-Security-Policy: frame-ancestors 'none';"
)

doc.add_paragraph()

# ── Finding 5 ──
add_heading('Finding 5: Lucky Winner Endpoint Exposes Prize State', 2, color=(160,100,0))
p = doc.add_paragraph()
r = p.add_run('Severity: MEDIUM  |  Category: Information Disclosure  |  CVSS: 4.8')
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RGBColor(160, 100, 0)

add_body('Where: /api/lucky-winner (fully public, unrated)', italic=True)
add_body('Response observed:')
add_code('{"wallet":"","amount":0,"txid":"","week":0,"live":false}')
add_body(
    'This endpoint is publicly accessible with no rate limiting. An attacker can poll it to monitor '
    'exactly when "live: true" fires, enabling precise timing of score manipulation attacks during '
    'prize distribution windows. The week numbering also enables historical winner enumeration.'
)
add_body('Remediation: Only return winner data when live:true. Rate-limit to 1 req/min per IP. Move winner selection entirely server-side.')

doc.add_paragraph()

# ── Finding 6 ──
add_heading('Finding 6: Missing HTTP Strict Transport Security', 2, color=(0,120,0))
p = doc.add_paragraph()
r = p.add_run('Severity: LOW  |  Category: Transport Security  |  CVSS: 3.7')
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RGBColor(0, 120, 0)

add_body('Missing: Strict-Transport-Security header on all responses.', italic=True)
add_body('Enables SSL stripping attacks on users without HSTS cached in their browser. In captive portals or public WiFi, first-time visitors may be served over HTTP.')
add_body('Remediation:')
add_code('Strict-Transport-Security: max-age=31536000; includeSubDomains; preload')

doc.add_paragraph()

# ── Finding 7 ──
add_heading('Finding 7: Treasury Wallet & Token Contract Hardcoded in Client', 2, color=(0,120,0))
p = doc.add_paragraph()
r = p.add_run('Severity: LOW  |  Category: Information Disclosure  |  CVSS: 3.1')
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RGBColor(0, 120, 0)

add_body('Values exposed in client JavaScript bundle:', italic=True)
add_code(
    "Treasury Wallet: CgfLiCwqpVbUX9tCvHe1chmWpEMuqzpcPSPkdwGtjrWv\n"
    "$DUMB Token:     8SCLpKVehUzuFBkqnF1qZu5LrodQDRCEoPVpCsrspump"
)
add_body('While on-chain data is public, embedding these in the client bundle makes it trivial to build convincing phishing clones using the real contract addresses. Also enables real-time treasury monitoring to time attacks.')
add_body('Remediation: Move to server-side env vars. Only expose treasury address when building transaction client-side. Add Solana Explorer deeplinks for legitimacy verification.')

doc.add_paragraph()

# ── Finding 8 ──
add_heading('Finding 8: Leaderboard 503 Under Load', 2, color=(0,120,0))
p = doc.add_paragraph()
r = p.add_run('Severity: LOW  |  Category: Reliability / Race Condition  |  CVSS: 3.5')
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RGBColor(0, 120, 0)

add_body('Where: /api/leaderboard — intermittent 503 Service Unavailable observed during assessment.', italic=True)
add_body('If the leaderboard endpoint is down during a weekly scoring cutoff, winner selection may run against stale data or silently fail. This creates an unpredictable outcome for prize distribution — and potentially a window for manipulation if an attacker can induce 503 conditions.')
add_body('Remediation: Add Vercel KV caching with 60s TTL + stale-while-revalidate. Set up uptime monitoring with PagerDuty alerts for this endpoint specifically.')

doc.add_paragraph()

# ── Finding 9 ──
add_heading('Finding 9: Balance Endpoints Unauthenticated & Unrated', 2, color=(0,120,0))
p = doc.add_paragraph()
r = p.add_run('Severity: LOW / INFO  |  Category: Privacy / Rate Limiting  |  CVSS: 2.9')
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RGBColor(0, 120, 0)

add_body('Affected: /api/balance?wallet=, /api/credits/treasury-check?wallet=', italic=True)
add_body('Any wallet address can be queried with no authentication or rate limiting. Enables enumeration of all user credit balances and treasury state. Without rate limiting these can be scraped to build a full user dataset.')
add_body('Remediation: Rate-limit to 10 req/min per IP. Optionally require wallet signature to view own balance only.')

# ═══════════════════════════════════════
# HIGH-PROBABILITY ATTACK SCENARIO
# ═══════════════════════════════════════
doc.add_page_break()
add_heading('4. High-Probability Attack: Score Inflation to Win Prize Pool', 1)

add_body(
    'This is the highest-probability real-world attack given the current implementation. '
    'It requires no special tools — only browser DevTools and a valid wallet.'
)

add_body('Attack Flow:', bold=True)
add_code(
    "Step 1: Attacker plays one legitimate game session\n"
    "Step 2: DevTools Network tab → copy the /api/send-tx request body\n"
    "Step 3: Inspect the copied serialized transaction bytes\n"
    "Step 4: Decode memo data → modify score field to 99999\n"
    "Step 5: Re-sign with same wallet (attacker owns the keys)\n"
    "Step 6: Submit modified tx to /api/send-tx\n"
    "Step 7: If server only validates signature (not memo integrity) → leaderboard #1\n"
    "Step 8: Wait for weekly lucky winner selection → collect prize pool"
)

add_body(
    'The defense requires the server to validate a server-issued HMAC over the game result '
    'at submission time. If the HMAC is missing or uses a client-predictable secret, '
    'the attack succeeds. This finding requires immediate verification of the /api/send-tx '
    'server-side logic.'
)

# ═══════════════════════════════════════
# QUICK WIN REMEDIATION
# ═══════════════════════════════════════
add_heading('5. Quick Win: Vercel Security Headers (next.config.ts)', 1)

add_body('All header-based findings (3, 4, 6, partial 1) can be fixed in a single shipping commit:')

add_code(
    "const securityHeaders = [\n"
    "  { key: 'X-Frame-Options', value: 'DENY' },\n"
    "  { key: 'X-Content-Type-Options', value: 'nosniff' },\n"
    "  { key: 'Referrer-Policy', value: 'strict-origin-when-cross-origin' },\n"
    "  { key: 'Strict-Transport-Security',\n"
    "    value: 'max-age=31536000; includeSubDomains' },\n"
    "  { key: 'Permissions-Policy',\n"
    "    value: 'camera=(), microphone=(), geolocation=()' },\n"
    "  { key: 'Access-Control-Allow-Origin',\n"
    "    value: 'https://arcade.somethingdumb.xyz' },\n"
    "];\n\n"
    "const nextConfig = {\n"
    "  async headers() {\n"
    "    return [{ source: '/(.*)', headers: securityHeaders }];\n"
    "  },\n"
    "};"
)

# ═══════════════════════════════════════
# REMEDIATION PRIORITY TABLE
# ═══════════════════════════════════════
add_heading('6. Remediation Priority Matrix', 1)

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr2 = table2.rows[0].cells
for i, text in enumerate(['Priority', 'Fix', 'Effort', 'Impact']):
    r = hdr2[i].paragraphs[0].add_run(text)
    r.bold = True
    r.font.size = Pt(10)

priorities = [
    ('🔴 1', 'Server-side score HMAC validation in /api/send-tx', 'High', 'Critical — prevents prize pool theft'),
    ('🔴 2', 'Remove unsafe-inline / unsafe-eval from CSP', 'Medium', 'High — blocks XSS wallet drain'),
    ('🟡 3', 'Restrict CORS to same origin', 'Low', 'Medium — blocks cross-origin probing'),
    ('🟡 4', 'Add X-Frame-Options: DENY', 'Low — 1 line', 'Medium — prevents clickjacking'),
    ('🟡 5', 'Rate-limit /api/lucky-winner, hide pre-winner state', 'Low', 'Medium — prevents timing attacks'),
    ('🟢 6', 'Add HSTS header', 'Low — 1 line', 'Low — hardens transport'),
    ('🟢 7', 'Cache leaderboard to prevent 503s', 'Medium', 'Low-Med — closes race window'),
    ('🟢 8', 'Rate-limit balance / treasury endpoints', 'Low', 'Low — prevents enumeration'),
]

for pri, fix, effort, impact in priorities:
    row = table2.add_row().cells
    row[0].paragraphs[0].add_run(pri).font.size = Pt(9)
    row[1].paragraphs[0].add_run(fix).font.size = Pt(9)
    row[2].paragraphs[0].add_run(effort).font.size = Pt(9)
    row[3].paragraphs[0].add_run(impact).font.size = Pt(9)

doc.add_paragraph()

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('This report is confidential — prepared for authorized security review only.')
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(120, 120, 120)

# ── Save ──
output_path = '/Users/hschaheen/Documents/Pentagonal/arcade_security_audit.docx'
doc.save(output_path)
print(f'✅ Saved: {output_path}')
